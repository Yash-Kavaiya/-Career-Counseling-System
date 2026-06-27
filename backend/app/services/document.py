"""Extract resume content from uploaded files (zero-infra, no system deps).

PDFs/text → extracted text via PyMuPDF; if a PDF yields little text (scanned),
its pages are rasterized to images for the vision model. DOCX via python-docx.
Images are passed straight through to the vision model.
"""
from __future__ import annotations

import base64
import io
import logging
from dataclasses import dataclass, field

import fitz  # PyMuPDF

logger = logging.getLogger(__name__)

MAX_VISION_PAGES = 3
_MIN_TEXT_CHARS = 80  # below this a PDF is treated as scanned


@dataclass
class ResumeContent:
    text: str = ""
    images: list[bytes] = field(default_factory=list)  # PNG bytes
    source: str = "text"  # text | vision

    def image_data_urls(self) -> list[str]:
        return [
            f"data:image/png;base64,{base64.b64encode(img).decode()}"
            for img in self.images
        ]


def _pdf_text(data: bytes) -> tuple[str, list[bytes]]:
    text_parts: list[str] = []
    images: list[bytes] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text())
        text = "\n".join(text_parts).strip()
        if len(text) < _MIN_TEXT_CHARS:  # scanned → rasterize for vision
            for page in doc[:MAX_VISION_PAGES]:
                pix = page.get_pixmap(dpi=150)
                images.append(pix.tobytes("png"))
    return text, images


def _docx_text(data: bytes) -> str:
    from docx import Document  # imported lazily

    document = Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_resume(filename: str, data: bytes) -> ResumeContent:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    try:
        if ext == "pdf":
            text, images = _pdf_text(data)
            return ResumeContent(
                text=text, images=images, source="vision" if images else "text"
            )
        if ext == "docx":
            return ResumeContent(text=_docx_text(data))
        if ext in {"txt", "md", "rtf"}:
            return ResumeContent(text=data.decode("utf-8", errors="ignore"))
        if ext in {"png", "jpg", "jpeg", "webp"}:
            return ResumeContent(images=[data], source="vision")
    except Exception:  # noqa: BLE001
        logger.exception("Failed to extract %s; falling back to raw decode", filename)
    # Unknown/failed → best-effort raw text
    return ResumeContent(text=data.decode("utf-8", errors="ignore"))
